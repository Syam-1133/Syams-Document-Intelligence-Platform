"""
Document loading utilities
"""
import os
import tempfile
import streamlit as st
from langchain_community.document_loaders import PyPDFDirectoryLoader, PyPDFLoader, TextLoader

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

def load_docx_file(file_path):
    """Load DOCX file using python-docx"""
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        if full_text:
            from langchain_core.documents import Document as LangchainDocument
            content = "\n".join(full_text)
            return [LangchainDocument(page_content=content, metadata={"source": file_path})]
        else:
            return []
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return []

def load_uploaded_documents(uploaded_files):
    """Load documents from uploaded files"""
    documents = []
    for uploaded_file in uploaded_files:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            # Load based on file type
            if uploaded_file.name.lower().endswith('.pdf'):
                loader = PyPDFLoader(tmp_file_path)
                docs = loader.load()
                documents.extend(docs)
                st.success(f"✅ Loaded PDF: {uploaded_file.name} ({len(docs)} pages)")
                
            elif uploaded_file.name.lower().endswith('.txt'):
                loader = TextLoader(tmp_file_path)
                docs = loader.load()
                documents.extend(docs)
                st.success(f"✅ Loaded TXT: {uploaded_file.name}")
                
            elif uploaded_file.name.lower().endswith(('.doc', '.docx')):
                if DOCX_AVAILABLE:
                    docs = load_docx_file(tmp_file_path)
                    if docs:
                        documents.extend(docs)
                        st.success(f"✅ Loaded DOCX: {uploaded_file.name}")
                    else:
                        st.warning(f"⚠️ No readable content found in: {uploaded_file.name}")
                else:
                    st.warning(f"⚠️ DOCX support not available. Install python-docx: pip install python-docx")
                
            else:
                st.warning(f"⚠️ Unsupported file type: {uploaded_file.name}")
                
        except Exception as e:
            st.error(f"❌ Error loading {uploaded_file.name}: {str(e)}")
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
    
    return documents

def load_documents_from_directory():
    """Load and cache documents from research_papers directory"""
    try:
        if os.path.exists("research_papers"):
            # Check for PDF files only (ignore hidden files like .DS_Store)
            pdf_files = [f for f in os.listdir("research_papers") if f.endswith('.pdf') and not f.startswith('.')]
            if pdf_files:
                loader = PyPDFDirectoryLoader("research_papers")
                docs = loader.load()
                if docs:
                    st.info(f"✅ Found {len(pdf_files)} PDF files in research_papers directory")
                return docs
        return []
    except Exception as e:
        st.error(f"Error loading documents: {str(e)}")
        return []
