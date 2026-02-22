import os
import tempfile
import shutil
# Apply OpenMP fix BEFORE importing other libraries
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'
os.environ['OMP_NUM_THREADS'] = '1'

# Configure page with dark theme FIRST
import streamlit as st
st.set_page_config(
    page_title="Document Intelligence Platform",
    page_icon="🔍",
    layout="wide",
    initial_sidebar_state="expanded"
)

from langchain_groq import ChatGroq
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import PyPDFDirectoryLoader, PyPDFLoader
from langchain_community.document_loaders import TextLoader
import time
import pandas as pd
from datetime import datetime

# Voice assistant imports
import pyttsx3
import speech_recognition as sr
import threading
import queue
import io
from gtts import gTTS
import base64
import requests
import json

# Alternative DOCX loader
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    st.warning("python-docx not available. Install with: pip install python-docx")

from dotenv import load_dotenv
load_dotenv()

# Dark theme CSS with professional styling - Enhanced for voice features
st.markdown("""
<style>
    .main {
        background-color: #0e1117;
        color: #ffffff;
    }
    .main-header {
        font-size: 2.8rem;
        color: #00d4aa;
        margin-bottom: 1rem;
        font-weight: 700;
        text-align: center;
        background: linear-gradient(135deg, #00d4aa, #0099cc);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-shadow: 0px 2px 10px rgba(0, 212, 170, 0.3);
    }
    .sub-header {
        font-size: 1.4rem;
        color: #66fcf1;
        margin-bottom: 1.5rem;
        text-align: center;
        font-weight: 300;
    }
    .metric-card {
        background: linear-gradient(135deg, #1f2833, #0b0c10);
        padding: 1.5rem;
        border-radius: 12px;
        border-left: 4px solid #00d4aa;
        margin-bottom: 1rem;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
        border: 1px solid #45a29e;
    }
    .success-box {
        background: linear-gradient(135deg, #1a472a, #0d1f14);
        border: 1px solid #2e8b57;
        border-radius: 8px;
        padding: 1.2rem;
        margin: 1rem 0;
        color: #90ee90;
    }
    .warning-box {
        background: linear-gradient(135deg, #4a3c1a, #2d240f);
        border: 1px solid #ffd700;
        border-radius: 8px;
        padding: 1.2rem;
        margin: 1rem 0;
        color: #ffd700;
    }
    .response-box {
        background: rgba(15, 23, 42, 0.95);
        border: 2px solid;
        border-image: linear-gradient(135deg, #00f5ff, #0066ff, #9d4edd) 1;
        border-radius: 8px;
        padding: 2rem;
        margin: 1rem 0;
        box-shadow: 0 8px 32px rgba(0, 102, 255, 0.4),
                0 0 25px rgba(0, 245, 255, 0.3);
        color: #e2e8f0;
        line-height: 1.8;
        min-height: 300px;
        width: 100%;
        max-width: 100%;
        overflow-y: auto;
    }
    .source-box {
        background: linear-gradient(135deg, #2d3047, #1a1c2b);
        border: 1px solid #5d5f7e;
        border-radius: 6px;
        padding: 1.2rem;
        margin: 0.8rem 0;
        color: #c5c6c7;
    }
    .sidebar .sidebar-content {
        background: linear-gradient(180deg, #1f2833, #0b0c10);
        color: white;
    }
    .stButton button {
        background: linear-gradient(135deg, #00d4aa, #0099cc);
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    .stButton button:hover {
        background: linear-gradient(135deg, #0099cc, #00d4aa);
        transform: translateY(-2px);
        box-shadow: 0 4px 15px rgba(0, 212, 170, 0.4);
    }
    .stTextInput input, .stTextArea textarea {
        background: #1f2833 !important;
        color: white !important;
        border: 1px solid #45a29e !important;
        border-radius: 8px !important;
    }
    .stExpander {
        background: #1f2833;
        border: 1px solid #45a29e;
        border-radius: 8px;
    }
    .tab-content {
        background: #0e1117;
        padding: 1rem;
        border-radius: 10px;
    }
    .status-indicator {
        display: inline-block;
        width: 10px;
        height: 10px;
        border-radius: 50%;
        margin-right: 8px;
    }
    .status-online {
        background-color: #00ff00;
        box-shadow: 0 0 10px #00ff00;
    }
    .status-offline {
        background-color: #ff4444;
        box-shadow: 0 0 10px #ff4444;
    }
    .chat-bubble {
        background: linear-gradient(135deg, #2d3047, #1a1c2b);
        border-radius: 15px;
        padding: 1rem;
        margin: 0.5rem 0;
        border: 1px solid #5d5f7e;
    }
    .chat-question {
        background: linear-gradient(135deg, #1a472a, #0d1f14);
        border-left: 4px solid #00d4aa;
    }
    .chat-answer {
        background: linear-gradient(135deg, #1a1f2e, #0f131f);
        border-left: 4px solid #4cc9f0;
    }
    .upload-section {
        background: linear-gradient(135deg, #1f2833, #0b0c10);
        padding: 2rem;
        border-radius: 12px;
        border: 2px dashed #45a29e;
        text-align: center;
        margin-bottom: 2rem;
    }
    .voice-controls {
        background: linear-gradient(135deg, #2d1a4c, #1a1c2b);
        padding: 1.5rem;
        border-radius: 12px;
        border: 1px solid #8a2be2;
        margin: 1rem 0;
    }
    .voice-active {
        background: linear-gradient(135deg, #1a472a, #0d1f14);
        border: 2px solid #00ff00;
        animation: pulse 2s infinite;
    }
    .premium-voice {
        background: linear-gradient(135deg, #4a1a6c, #2d1a4c);
        border: 2px solid #ff6b6b;
        padding: 1.5rem;
        border-radius: 12px;
        margin: 1rem 0;
    }
    @keyframes pulse {
        0% { border-color: #00ff00; }
        50% { border-color: #00d4aa; }
        100% { border-color: #00ff00; }
    }
    .voice-selection {
        background: linear-gradient(135deg, #1f2833, #0b0c10);
        padding: 1rem;
        border-radius: 10px;
        border: 1px solid #45a29e;
        margin: 0.5rem 0;
    }
    .voice-option {
        padding: 0.8rem;
        margin: 0.3rem 0;
        border-radius: 8px;
        cursor: pointer;
        transition: all 0.3s ease;
        border: 1px solid #45a29e;
    }
    .voice-option:hover {
        background: linear-gradient(135deg, #2d3047, #1a1c2b);
        border-color: #8a2be2;
    }
    .voice-option.selected {
        background: linear-gradient(135deg, #4a1a6c, #2d1a4c);
        border-color: #ff6b6b;
    }
</style>
""", unsafe_allow_html=True)

# Configure API keys
os.environ['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')
os.environ['GROQ_API_KEY'] = os.getenv('GROQ_API_KEY')
ELEVENLABS_API_KEY = os.getenv('ELEVENLABS_API_KEY')

groq_api_key = os.getenv('GROQ_API_KEY')

# Validate API keys
if not groq_api_key:
    st.error("🔑 GROQ API key not found! Please check your .env file.")
    st.stop()

if not os.getenv('OPENAI_API_KEY'):
    st.error("🔑 OpenAI API key not found! Please check your .env file.")
    st.stop()

# Initialize voice assistant components
@st.cache_resource
def init_voice_engine():
    """Initialize text-to-speech engine"""
    try:
        engine = pyttsx3.init()
        voices = engine.getProperty('voices')
        if voices:
            engine.setProperty('voice', voices[1].id)  # Female voice if available
        engine.setProperty('rate', 150)
        engine.setProperty('volume', 0.8)
        return engine
    except Exception as e:
        st.warning(f"Text-to-speech engine initialization failed: {str(e)}")
        return None

@st.cache_resource
def init_speech_recognizer():
    """Initialize speech recognition"""
    return sr.Recognizer()

# Initialize voice components
tts_engine = init_voice_engine()
speech_recognizer = init_speech_recognizer()

# ElevenLabs Voice Options
ELEVENLABS_VOICES = {
    "Rachel": "21m00Tcm4TlvDq8ikWAM",
    "Domi": "AZnzlk1XvdvUeBnXmlld",
    "Bella": "EXAVITQu4vr4xnSDxMaL",
    "Antoni": "ErXwobaYiN019PkySvjV",
    "Elli": "MF3mGyEYCl7XYWbV9V6O",
    "Josh": "TxGEqnHWrfWFTfGW9XjX",
    "Arnold": "VR6AewLTigWG4xSOukaG",
    "Adam": "pNInz6obpgDQGcFmaJgB",
    "Sam": "yoZ06aMxZJJ28mfd3POQ"
}

# Initialize session state for voice assistant
if "listening" not in st.session_state:
    st.session_state.listening = False
if "audio_queue" not in st.session_state:
    st.session_state.audio_queue = queue.Queue()
if "last_spoken_text" not in st.session_state:
    st.session_state.last_spoken_text = ""
if "selected_voice" not in st.session_state:
    st.session_state.selected_voice = "Rachel"
if "tts_provider" not in st.session_state:
    st.session_state.tts_provider = "elevenlabs"

def load_docx_file(file_path):
    """Load DOCX file using python-docx"""
    try:
        doc = Document(file_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        if full_text:
            from langchain_core.documents import Document as LangchainDocument
            content = "\n".join(full_text)
            return [LangchainDocument(page_content=content, metadata={"source": file_path})]
        else:
            return []
    except Exception as e:
        st.error(f"Error reading DOCX file: {str(e)}")
        return []

def text_to_speech_elevenlabs(text, voice_id, stability=0.5, similarity_boost=0.5):
    """Convert text to speech using ElevenLabs API"""
    try:
        if not ELEVENLABS_API_KEY:
            return "❌ ElevenLabs API key not configured"
        
        # ElevenLabs API endpoint
        url = f"https://api.elevenlabs.io/v1/text-to-speech/{ELEVENLABS_VOICES.get(voice_id, '21m00Tcm4TlvDq8ikWAM')}"
        
        headers = {
            "Accept": "audio/mpeg",
            "Content-Type": "application/json",
            "xi-api-key": ELEVENLABS_API_KEY
        }
        
        data = {
            "text": text,
            "model_id": "eleven_turbo_v2_5",
            "voice_settings": {
                "stability": stability,
                "similarity_boost": similarity_boost
            }
        }
        
        response = requests.post(url, json=data, headers=headers)
        
        if response.status_code == 200:
            audio_buffer = io.BytesIO(response.content)
            audio_buffer.seek(0)
            
            # Encode audio for HTML audio element
            audio_base64 = base64.b64encode(audio_buffer.read()).decode()
            audio_html = f'''
                <audio autoplay controls style="width: 100%; margin: 10px 0;">
                    <source src="data:audio/mp3;base64,{audio_base64}" type="audio/mp3">
                    Your browser does not support the audio element.
                </audio>
                <div style="text-align: center; color: #ff6b6b; font-size: 0.9rem;">
                    🎙️ Premium Voice: {list(ELEVENLABS_VOICES.keys())[list(ELEVENLABS_VOICES.values()).index(voice_id)]}
                </div>
            '''
            return audio_html
        else:
            return f"❌ ElevenLabs API Error: {response.status_code} - {response.text}"
            
    except Exception as e:
        return f"❌ Error in ElevenLabs TTS: {str(e)}"

def text_to_speech(text, provider="elevenlabs", voice_id=None):
    """Convert text to speech using selected provider"""
    try:
        if provider == "elevenlabs" and ELEVENLABS_API_KEY:
            if not voice_id:
                voice_id = ELEVENLABS_VOICES.get(st.session_state.selected_voice, "21m00Tcm4TlvDq8ikWAM")
            return text_to_speech_elevenlabs(text[:800], voice_id)  # Limit text length
                
    except Exception as e:
        return f"❌ Error in text-to-speech: {str(e)}"

def speech_to_text():
    """Convert speech to text using microphone input"""
    try:
        with sr.Microphone() as source:
            st.session_state.listening = True
            speech_recognizer.adjust_for_ambient_noise(source, duration=0.5)
            
            # Show listening indicator
            listening_placeholder = st.empty()
            listening_placeholder.info("🎤 Listening... Speak now!")
            
            # Listen for audio
            audio = speech_recognizer.listen(source, timeout=10, phrase_time_limit=15)
            
            # Recognize speech
            text = speech_recognizer.recognize_google(audio)
            listening_placeholder.empty()
            st.session_state.listening = False
            return text
            
    except sr.WaitTimeoutError:
        st.session_state.listening = False
        return "Timeout: No speech detected"
    except sr.UnknownValueError:
        st.session_state.listening = False
        return "Could not understand audio"
    except sr.RequestError as e:
        st.session_state.listening = False
        return f"Error with speech recognition service: {str(e)}"
    except Exception as e:
        st.session_state.listening = False
        return f"Error: {str(e)}"

def voice_query_interface():
    """Voice query interface component"""

    with st.container():
        col1, col2, col3, col4 = st.columns([1, 1, 1, 1])
        
        with col1:
            if st.button("🎤 Start Listening", use_container_width=True, type="primary"):
                if "vectors" not in st.session_state:
                    st.error("Please process documents first!")
                else:
                    spoken_text = speech_to_text()
                    if spoken_text and not spoken_text.startswith(("Timeout", "Could not", "Error")):
                        st.session_state.voice_query = spoken_text
                        st.success(f"🎤 Heard: '{spoken_text}'")
                        
                        # Process the voice query
                        response, response_time = process_query(spoken_text)
                        if response:
                            # Store in chat history
                            st.session_state.chat_history.append({
                                "question": spoken_text,
                                "answer": response['answer'],
                                "timestamp": datetime.now(),
                                "response_time": response_time,
                                "voice_query": True
                            })
                            
                            # Store for voice replay
                            st.session_state.last_spoken_text = response['answer']
                            st.session_state.show_response = True
                            st.session_state.current_response = response['answer']
                    else:
                        st.error(f"Voice recognition failed: {spoken_text}")
        
        with col2:
            if st.session_state.get('last_spoken_text'):
                if st.button("🔊 Repeat Answer", use_container_width=True):
                    audio_response = text_to_speech(
                        st.session_state.last_spoken_text,
                        provider=st.session_state.tts_provider,
                        voice_id=ELEVENLABS_VOICES[st.session_state.selected_voice]
                    )
                    if audio_response:
                        st.markdown(audio_response, unsafe_allow_html=True)
        
        with col3:
            # Voice preview
            if st.button("🎵 Preview Voice", use_container_width=True):
                preview_text = "Hello! I'm your Shyam AI voice assistant. How can I help you today?"
                audio_response = text_to_speech(
                    preview_text,
                    provider=st.session_state.tts_provider,
                    voice_id=ELEVENLABS_VOICES[st.session_state.selected_voice]
                )
                if audio_response:
                    st.markdown(audio_response, unsafe_allow_html=True)
        
        with col4:
            if st.button("⏹️ Stop Audio", use_container_width=True):
                st.info("Refresh page to stop audio playback")
    
    # Display response outside columns for full width
    if st.session_state.get('show_response') and st.session_state.get('current_response'):
        st.markdown("### 📋 AI RESPONSE")
        st.markdown(f'<div class="response-box">{st.session_state.current_response}</div>', unsafe_allow_html=True)
        
        # Text-to-speech for the response
        st.markdown("### 🔊 AI AUDIO RESPONSE")
        audio_response = text_to_speech(
            st.session_state.current_response, 
            provider=st.session_state.tts_provider,
            voice_id=ELEVENLABS_VOICES[st.session_state.selected_voice]
        )
        if audio_response:
            st.markdown(audio_response, unsafe_allow_html=True)
        
        # Reset the flag after displaying
        st.session_state.show_response = False

# Initialize LLM with error handling
@st.cache_resource
def get_llm():
    try:
        return ChatGroq(groq_api_key=groq_api_key, model="llama-3.1-8b-instant", temperature=0.1)
    except Exception as e:
        st.error(f"Error initializing Groq LLM: {str(e)}")
        return None

llm = get_llm()
if llm is None:
    st.stop()

# Create enhanced prompt template
prompt = ChatPromptTemplate.from_template(
    """
    You are an expert in analyzing documents please talk in a conversational tone and like friendly chatbot. 
    Provide comprehensive, accurate answers based strictly on the provided context.
    
    GUIDELINES:
    - Answer the question using only the information from the provided context
    - If the information is not in the context, clearly state this
    - Provide detailed explanations when appropriate
    - Include relevant technical details from the research papers
    - Structure your response clearly and professionally
    - please maintain a conversational tone with the user queries
    
    CONTEXT: {context}
    
    QUESTION: {input}
    
    Please provide a thorough, well-structured answer:
    """
)

def load_uploaded_documents(uploaded_files):
    """Load documents from uploaded files"""
    documents = []
    for uploaded_file in uploaded_files:
        # Create a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            # Load based on file type
            if uploaded_file.name.lower().endswith('.pdf'):
                loader = PyPDFLoader(tmp_file_path)
                docs = loader.load()
                documents.extend(docs)
                st.success(f"✅ Loaded PDF: {uploaded_file.name} ({len(docs)} pages)")
                
            elif uploaded_file.name.lower().endswith('.txt'):
                loader = TextLoader(tmp_file_path)
                docs = loader.load()
                documents.extend(docs)
                st.success(f"✅ Loaded TXT: {uploaded_file.name}")
                
            elif uploaded_file.name.lower().endswith(('.doc', '.docx')):
                if DOCX_AVAILABLE:
                    docs = load_docx_file(tmp_file_path)
                    if docs:
                        documents.extend(docs)
                        st.success(f"✅ Loaded DOCX: {uploaded_file.name}")
                    else:
                        st.warning(f"⚠️ No readable content found in: {uploaded_file.name}")
                else:
                    st.warning(f"⚠️ DOCX support not available. Install python-docx: pip install python-docx")
                
            else:
                st.warning(f"⚠️ Unsupported file type: {uploaded_file.name}")
                
        except Exception as e:
            st.error(f"❌ Error loading {uploaded_file.name}: {str(e)}")
        finally:
            # Clean up temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
    
    return documents

@st.cache_data
def load_documents():
    """Load and cache documents from research_papers directory"""
    try:
        if os.path.exists("research_papers"):
            # Check for PDF files only (ignore hidden files like .DS_Store)
            pdf_files = [f for f in os.listdir("research_papers") if f.endswith('.pdf') and not f.startswith('.')]
            if pdf_files:
                loader = PyPDFDirectoryLoader("research_papers")
                docs = loader.load()
                if docs:
                    st.info(f"✅ Found {len(pdf_files)} PDF files in research_papers directory")
                return docs
        return []
    except Exception as e:
        st.error(f"Error loading documents: {str(e)}")
        return []

def create_vector_embedding_from_files(uploaded_files=None):
    """Create vector embeddings from uploaded files or existing documents"""
    try:
        # Create progress containers
        progress_container = st.container()
        status_container = st.container()
        metrics_container = st.container()
        
        with progress_container:
            st.markdown("### 📊 Document Processing Progress")
            progress_bar = st.progress(0)
            status_text = st.empty()
            time_elapsed = st.empty()
            
        start_time = time.time()
        
        # Step 1: Load documents
        status_text.text("📄 Loading documents...")
        progress_bar.progress(10)
        
        if uploaded_files:
            # Use uploaded files
            docs = load_uploaded_documents(uploaded_files)
            source_type = "uploaded files"
        else:
            # Use existing research_papers directory
            docs = load_documents()
            source_type = "research_papers directory"
        
        if not docs:
            st.error("❌ No documents found or could be loaded!")
            return False
        
        st.info(f"📑 Loaded {len(docs)} documents from {source_type}")
        
        # Step 2: Initialize embeddings
        status_text.text("🔧 Initializing OpenAI embeddings...")
        progress_bar.progress(30)
        
        embeddings = OpenAIEmbeddings(
            chunk_size=1000,
            max_retries=3
        )
        
        # Step 3: Split documents
        status_text.text("✂️ Processing and splitting documents into chunks...")
        progress_bar.progress(50)
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        final_documents = text_splitter.split_documents(docs)
        
        if not final_documents:
            st.error("❌ No text could be extracted from the documents!")
            return False
        
        st.info(f"📑 Created {len(final_documents)} text chunks")
        
        # Step 4: Create vector store
        status_text.text("🔍 Creating vector database (this may take a moment)...")
        progress_bar.progress(70)
        
        # Process in batches
        batch_size = 10
        vectors = None
        
        for i in range(0, len(final_documents), batch_size):
            batch = final_documents[i:i+batch_size]
            if vectors is None:
                vectors = FAISS.from_documents(batch, embeddings)
            else:
                batch_vectors = FAISS.from_documents(batch, embeddings)
                vectors.merge_from(batch_vectors)
            
            progress = 70 + (25 * (i + batch_size) / len(final_documents))
            progress_bar.progress(min(95, int(progress)))
            
            # Update time
            elapsed = time.time() - start_time
            time_elapsed.text(f"⏱️ Time elapsed: {elapsed:.1f} seconds")
        
        st.session_state.vectors = vectors
        st.session_state.final_documents = final_documents
        st.session_state.document_source = source_type
        
        progress_bar.progress(100)
        status_text.text("✅ Vector database created successfully!")
        
        total_time = time.time() - start_time
        
        # Display metrics
        with metrics_container:
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Documents Processed", len(docs))
            with col2:
                st.metric("Text Chunks Created", len(final_documents))
            with col3:
                st.metric("Processing Time", f"{total_time:.1f}s")
        
        return True
        
    except Exception as e:
        st.error(f"❌ Error creating vector embeddings: {str(e)}")
        return False

def process_query(user_prompt):
    """Process user query safely"""
    try:
        with st.spinner("🔍 Searching through documents and generating response..."):
            # Create retriever
            retriever = st.session_state.vectors.as_retriever(
                search_type="similarity",
                search_kwargs={"k": 4}
            )
            
            # Get relevant documents
            docs = retriever.invoke(user_prompt)
            
            # Format context
            context = "\n\n".join([doc.page_content for doc in docs])
            
            # Create chain with LCEL
            chain = (
                {"context": lambda x: context, "input": RunnablePassthrough()}
                | prompt
                | llm
                | StrOutputParser()
            )
            
            # Process query
            start_time = time.time()
            answer = chain.invoke(user_prompt)
            end_time = time.time()
            
            # Return in expected format
            response = {
                "answer": answer,
                "context": docs
            }
            
            return response, end_time - start_time
            
    except Exception as e:
        st.error(f"❌ Error processing query: {str(e)}")
        return None, 0

# Initialize session state for voice assistant
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "document_source" not in st.session_state:
    st.session_state.document_source = "Not initialized"
if "voice_query" not in st.session_state:
    st.session_state.voice_query = ""

# Main UI - Professional Header
st.markdown('<div class="main-header">Syam\'s AI-Powered Document Analyzer</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="sub-header">'
    '🚀 AI-Powered Document Analysis & Knowledge Extraction System<br>'
    '<small style="color: #88d3ce; font-size: 1.1rem; font-weight: 300;">'
    'Upload documents • Ask questions • Get answers • Voice Assistant Bot'
    '</small>'
    '</div>', 
    unsafe_allow_html=True
)

# Sidebar with enhanced professional layout
with st.sidebar:
    st.markdown("### 🏢 SYSTEM DASHBOARD")
    
    # System Status
    st.markdown("#### 📈 SYSTEM STATUS")
    if "vectors" in st.session_state:
        st.markdown(
            f'<div class="success-box">'
            f'<span class="status-indicator status-online"></span>'
            f'<strong>VECTOR DATABASE: ONLINE</strong><br>'
            f'Source: {st.session_state.document_source}'
            f'</div>', 
            unsafe_allow_html=True
        )
        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f'<div class="metric-card">📄 Document Chunks<br><h3>{len(st.session_state.final_documents)}</h3></div>', unsafe_allow_html=True)
        with col2:
            st.markdown(f'<div class="metric-card">🤖 Model<br><h3>Llama-3.1-8B</h3></div>', unsafe_allow_html=True)
    else:
        st.markdown(
            f'<div class="warning-box">'
            f'<span class="status-indicator status-offline"></span>'
            f'<strong>VECTOR DATABASE: OFFLINE</strong><br>'
            f'Upload documents or use existing ones'
            f'</div>', 
            unsafe_allow_html=True
        )
    
    st.markdown("---")
    
    # Quick Actions
    st.markdown("#### ⚡ QUICK ACTIONS")
    
    # Document Upload Section
    st.markdown("#### 📤 UPLOAD DOCUMENTS")
    uploaded_files = st.file_uploader(
        "Choose files",
        type=['pdf', 'txt', 'doc', 'docx'],
        accept_multiple_files=True,
        label_visibility="collapsed"
    )
    
    if uploaded_files:
        if st.button("🚀 PROCESS UPLOADED FILES", use_container_width=True, type="primary"):
            if create_vector_embedding_from_files(uploaded_files):
                st.success("Uploaded files processed successfully!")
                st.rerun()
    
    # Use existing files button
    if st.button("USE EXISTING DOCUMENTS", use_container_width=True):
        if create_vector_embedding_from_files():
            st.success("Existing documents processed successfully!")
            st.rerun()
    
    if st.button("CLEAR CHAT HISTORY", use_container_width=True):
        st.session_state.chat_history = []
        st.success("Chat history cleared!")
        st.rerun()
    
    st.markdown("---")
    
    st.markdown("#### 💬 RECENT QUERIES")
    if st.session_state.chat_history:
        for i, chat in enumerate(st.session_state.chat_history[-5:]):
            with st.expander(f"Q: {chat['question'][:50]}..." if len(chat['question']) > 50 else f"Q: {chat['question']}", expanded=False):
                voice_indicator = "🎙️ " if chat.get('voice_query') else ""
                st.markdown(f'<div class="chat-bubble chat-question"><strong>Question:</strong> {voice_indicator}{chat["question"]}</div>', unsafe_allow_html=True)
                st.markdown(f'<div class="chat-bubble chat-answer"><strong>Answer:</strong> {chat["answer"][:150]}...</div>', unsafe_allow_html=True)
    else:
        st.info("No queries yet. Start a conversation!")
    
    st.markdown("---")
    st.markdown("**Built with:**")
    st.markdown("• Streamlit")
    st.markdown("• LangChain ")
    st.markdown("• Groq  LLM")
    st.markdown("• ElevenLabs ")
    st.markdown(f"*Session: {datetime.now().strftime('%Y-%m-%d %H:%M')}*")

# Main content area
tab1, tab2, tab3, tab4 = st.tabs(["🔍 QUERY DOCUMENTS", "🎙️ VOICE ASSISTANT", "📊 ANALYTICS", "⚙️ SETTINGS"])

with tab1:
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.markdown("### 💬 DOCUMENT QUERY INTERFACE")
        
        # System initialization status
        if "vectors" not in st.session_state:
            st.markdown(
                '<div class="upload-section">'
                '<h3>📁 Upload Documents to Get Started</h3>'
                '<p>Upload PDF, TXT, or DOCX files to analyze their content</p>'
                '<p><small>Or use the "Use Existing Documents" button if you have files in the research_papers directory</small></p>'
                '</div>', 
                unsafe_allow_html=True
            )
        else:
            # Query interface
            st.markdown("#### 📝 ASK A QUESTION")
            user_prompt = st.text_area(
                "Enter your question about the documents:",
                placeholder="e.g., What are the main findings in these documents? Explain the key concepts...",
                height=120,
                label_visibility="collapsed"
            )
            
            col1, col2, col3 = st.columns([2, 1, 1])
            with col1:
                ask_button = st.button("🔍 ANALYZE DOCUMENTS", type="primary", use_container_width=True)
            with col2:
                speak_button = st.button("🔊 SPEAK ANSWER", use_container_width=True)
            
            if ask_button and user_prompt:
                response, response_time = process_query(user_prompt)
                
                if response:
                    # Store in chat history
                    st.session_state.chat_history.append({
                        "question": user_prompt,
                        "answer": response['answer'],
                        "timestamp": datetime.now(),
                        "response_time": response_time
                    })
                    
                    # Display response in full-width container
                    with st.container():
                        st.markdown("### 📋 ANALYSIS RESULTS")
                        st.markdown(f'<div class="response-box">{response["answer"]}</div>', unsafe_allow_html=True)
                    
                    # Response metrics
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.markdown(f'<div class="metric-card">⏱️ Response Time<br><h3>{response_time:.2f}s</h3></div>', unsafe_allow_html=True)
                    with col2:
                        st.markdown(f'<div class="metric-card">📚 Sources Retrieved<br><h3>{len(response["context"])}</h3></div>', unsafe_allow_html=True)
                    with col3:
                        st.markdown(f'<div class="metric-card">🤖 Model Used<br><h3>Llama-3.1-8B</h3></div>', unsafe_allow_html=True)
                    
                    # Store for voice
                    st.session_state.last_spoken_text = response['answer']
                    
                    # Source documents
                    with st.expander("📚 VIEW SOURCE DOCUMENTS (4 most relevant)", expanded=False):
                        for i, doc in enumerate(response['context']):
                            st.markdown(f"**Document {i+1}**")
                            st.markdown(f'<div class="source-box">{doc.page_content}</div>', unsafe_allow_html=True)
                            st.markdown("---")
            
            if speak_button and user_prompt and "last_spoken_text" in st.session_state:
                audio_response = text_to_speech(
                    st.session_state.last_spoken_text,
                    provider=st.session_state.tts_provider,
                    voice_id=ELEVENLABS_VOICES[st.session_state.selected_voice]
                )
                if audio_response:
                    st.markdown(audio_response, unsafe_allow_html=True)

with tab2:
    st.markdown("### 🎙️ VOICE-ENABLED DOCUMENT ASSISTANT")
    
    if "vectors" not in st.session_state:
        st.warning("⚠️ Please process documents first to use the voice assistant!")
        if st.button("📁 Go to Document Processing", use_container_width=True):
            st.rerun()
    else:
        voice_query_interface()
        
        # Voice chat history
        if any(chat.get('voice_query') for chat in st.session_state.chat_history):
            st.markdown("### 🎙️ VOICE QUERY HISTORY")
            voice_chats = [chat for chat in st.session_state.chat_history if chat.get('voice_query')]
            for chat in voice_chats[-5:]:
                with st.expander(f"🎙️ {chat['question'][:60]}...", expanded=False):
                    st.markdown(f'<div class="chat-bubble chat-question"><strong>Voice Question:</strong> {chat["question"]}</div>', unsafe_allow_html=True)
                    st.markdown(f'<div class="chat-bubble chat-answer"><strong>AI Response:</strong> {chat["answer"]}</div>', unsafe_allow_html=True)
                    
                    if st.button(f"🔊 Play Response", key=f"play_{chat['timestamp']}"):
                        audio_response = text_to_speech(
                            chat['answer'],
                            provider=st.session_state.tts_provider,
                            voice_id=ELEVENLABS_VOICES[st.session_state.selected_voice]
                        )
                        if audio_response:
                            st.markdown(audio_response, unsafe_allow_html=True)

with tab3:
    st.markdown("### 📊 SYSTEM ANALYTICS")
    
    if "vectors" in st.session_state:
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 📈 DOCUMENT STATISTICS")
            st.markdown(f"""
            <div style='background: linear-gradient(135deg, #1f2833, #0b0c10); padding: 1.5rem; border-radius: 10px; border: 1px solid #45a29e;'>
                <div style='display: grid; grid-template-columns: 1fr 1fr; gap: 1rem;'>
                    <div style='text-align: center;'>
                        <h4 style='color: #66fcf1; margin: 0;'>Total Documents</h4>
                        <h3 style='color: #00d4aa; margin: 0.5rem 0;'>{len(st.session_state.final_documents)}</h3>
                    </div>
                    <div style='text-align: center;'>
                        <h4 style='color: #66fcf1; margin: 0;'>Text Chunks</h4>
                        <h3 style='color: #00d4aa; margin: 0.5rem 0;'>{len(st.session_state.final_documents)}</h3>
                    </div>
                    <div style='text-align: center;'>
                        <h4 style='color: #66fcf1; margin: 0;'>Document Source</h4>
                        <h3 style='color: #00d4aa; margin: 0.5rem 0; font-size: 0.9rem;'>{st.session_state.document_source}</h3>
                    </div>
                    <div style='text-align: center;'>
                        <h4 style='color: #66fcf1; margin: 0;'>Vector Dimension</h4>
                        <h3 style='color: #00d4aa; margin: 0.5rem 0;'>1536</h3>
                    </div>
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        with col2:
            st.markdown("#### ⚡ PERFORMANCE METRICS")
            if st.session_state.chat_history:
                avg_response_time = sum([chat['response_time'] for chat in st.session_state.chat_history]) / len(st.session_state.chat_history)
                voice_queries = len([chat for chat in st.session_state.chat_history if chat.get('voice_query')])
                st.markdown(f"""
                <div style='background: linear-gradient(135deg, #1f2833, #0b0c10); padding: 1.5rem; border-radius: 10px; border: 1px solid #45a29e;'>
                    <div style='display: grid; grid-template-columns: 1fr 1fr; gap: 1rem;'>
                        <div style='text-align: center;'>
                            <h4 style='color: #66fcf1; margin: 0;'>Total Queries</h4>
                            <h3 style='color: #00d4aa; margin: 0.5rem 0;'>{len(st.session_state.chat_history)}</h3>
                        </div>
                        <div style='text-align: center;'>
                            <h4 style='color: #66fcf1; margin: 0;'>Voice Queries</h4>
                            <h3 style='color: #00d4aa; margin: 0.5rem 0;'>{voice_queries}</h3>
                        </div>
                        <div style='text-align: center;'>
                            <h4 style='color: #66fcf1; margin: 0;'>Avg Response Time</h4>
                            <h3 style='color: #00d4aa; margin: 0.5rem 0;'>{avg_response_time:.2f}s</h3>
                        </div>
                        <div style='text-align: center;'>
                            <h4 style='color: #66fcf1; margin: 0;'>Session Start</h4>
                            <h3 style='color: #00d4aa; margin: 0.5rem 0;'>{datetime.now().strftime('%H:%M')}</h3>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
            else:
                st.info("No query data available yet")

with tab4:
    st.markdown("### ⚙️ SYSTEM CONFIGURATION")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🤖 MODEL SETTINGS")
        model_option = st.selectbox("LLM Model", ["llama-3.1-8b-instant", "llama-3.1-70b-versatile", "mixtral-8x7b"], index=0)
        temperature = st.slider("Temperature", 0.0, 1.0, 0.1, 0.1)
        max_tokens = st.number_input("Max Tokens", 100, 4000, 1000)
        
        st.markdown("#### 🔍 RETRIEVAL SETTINGS")
        chunk_count = st.slider("Number of Chunks to Retrieve", 1, 10, 4)
        search_type = st.selectbox("Search Type", ["similarity", "mmr", "similarity_score_threshold"])
    
    with col2:
        st.markdown("#### 📄 DOCUMENT PROCESSING")
        chunk_size = st.number_input("Chunk Size", 500, 2000, 1000)
        chunk_overlap = st.number_input("Chunk Overlap", 0, 500, 200)
        text_splitter = st.selectbox("Text Splitter", ["RecursiveCharacterTextSplitter", "CharacterTextSplitter"])
        
        st.markdown("#### 🎙️ VOICE SETTINGS")
        tts_engine_choice = st.selectbox("TTS Engine", ["ElevenLabs (Premium)", "gTTS (Online)", "pyttsx3 (Offline)"])
        speech_rate = st.slider("Speech Rate", 50, 300, 150)
        auto_play_voice = st.checkbox("Auto-play voice responses", value=True)

# Footer
st.markdown("---")
st.markdown(
    "<div style='text-align: center; color: #66fcf1; font-size: 0.9rem;'>"
    "🔍 Document Intelligence Platform • Built with Streamlit, LangChain, and Groq • "
    "ElevenLabs Voice Assistant • Enterprise Ready AI Solution"
    "</div>", 
    unsafe_allow_html=True
)